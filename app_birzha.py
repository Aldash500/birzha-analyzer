
import streamlit as st
import datetime
from docx import Document
import os

st.set_page_config(page_title="Биржевой Анализатор", layout="wide")
st.title("Биржевой Анализатор — отчёты по товарным биржам")

# Выбор периода
st.subheader("Выберите месяц и год для анализа")
month = st.selectbox(
    "Месяц",
    ["Январь", "Февраль", "Март", "Апрель", "Май", "Июнь",
     "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"]
)
year = st.selectbox("Год", list(reversed(range(2020, datetime.datetime.now().year + 1))))

# Кнопка запуска анализа
if st.button("Сформировать отчёт"):
    with st.spinner("Идёт сбор и анализ публикаций по биржам..."):

        # --- Раздел 1: Новости с utb.kz ---
        utb_articles = [
            {
                "title": "Торги углём на UTB достигли рекорда в октябре",
                "summary": "На бирже utb.kz в октябре прошло рекордное количество сделок по углю, превышающее 3 млрд тенге.",
                "url": "https://utb.kz/news/ugol-torgi-oktyabr"
            }
        ]

        # --- Раздел 2: Новости из других источников ---
        other_articles = [
            {
                "title": "Kazakhstan commodity exchange sees coal volume spike",
                "summary": "Total coal trade on Kazakhstan's exchange increased by 21% in October, driven by export demand.",
                "translation": "Общий объём торговли углём на казахстанской бирже вырос на 21% в октябре на фоне экспортного спроса.",
                "url": "https://example.com/coal-volume-oct"
            }
        ]

        # Создание Word-документа
        doc = Document()
        doc.add_heading(f"Биржевой анализ — {month} {year}", 0)
        doc.add_paragraph(f"Период: {month} {year}")

        # Раздел 1
        doc.add_heading("Раздел 1. Новости с сайта utb.kz", level=1)
        doc.add_paragraph(f"Количество публикаций: {len(utb_articles)}")
        for i, a in enumerate(utb_articles, 1):
            doc.add_paragraph(f"{i}. {a['title']}")
            doc.add_paragraph(a["summary"])
            doc.add_paragraph(f"Источник: {a['url']}")
            doc.add_paragraph("")
        doc.add_paragraph("Вывод: Биржа UTB демонстрирует активный рост торгов углём, что свидетельствует о повышении спроса на сырьё.")
        doc.add_paragraph("Рекомендация: Продолжить развитие прозрачных механизмов торгов и расширять аналитическую отчётность.")

        # Раздел 2
        doc.add_heading("Раздел 2. Публикации из других источников", level=1)
        doc.add_paragraph(f"Количество публикаций: {len(other_articles)}")
        for i, a in enumerate(other_articles, 1):
            doc.add_paragraph(f"{i}. {a['title']}")
            doc.add_paragraph(f"[EN] {a['summary']}")
            if "translation" in a:
                doc.add_paragraph(f"[RU] Перевод: {a['translation']}")
            doc.add_paragraph(f"Источник: {a['url']}")
            doc.add_paragraph("")
        doc.add_paragraph("Вывод: Международные источники фиксируют рост биржевой активности и экспортный спрос на уголь.")
        doc.add_paragraph("Рекомендация: Усилить внимание к экспортным рынкам и инвестировать в логистику.")

        # Сохранение и загрузка
        file_name = f"birzha_report_{month}_{year}.docx"
        file_path = os.path.join("./", file_name)
        doc.save(file_path)

        with open(file_path, "rb") as f:
            st.download_button(
                label="Скачать отчёт в Word",
                data=f,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        os.remove(file_path)
